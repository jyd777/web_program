from docx import Document
from docx.shared import Inches

document = Document()

# Title
document.add_heading('Climate Change and Residential Property', 0)

# Paragraph 1
document.add_paragraph(
    'THINK ABOUT the places vulnerable to climate change, and you might picture rice paddies in Bangladesh or low-lying islands in the Pacific. '
    'But another, more surprising answer ought to be your own house. About a tenth of the world\'s residential property by value is under threat from global warming—including many houses that are nowhere near the coast.'
)

document.add_paragraph(
    '提起那些易受气候变化影响的地方，脑海中或许会浮现孟加拉国的水稻田，亦或是太平洋中海拔较低的岛屿。'
    '但更为令人惊奇的是另一个画面，也就是你自己的房子。按价值评估，世界上约有十分之一的房产面临全球变暖的威胁——包括许多内陆房屋。'
)

# Paragraph 2
document.add_paragraph(
    'From tornadoes battering midwestern American suburbs to tennis-ball-size hailstones smashing the roofs of Italian villas, the severe weather '
    'brought about by greenhouse-gas emissions is shaking the foundations of the world\'s most important asset class.'
)

document.add_paragraph(
    '从肆虐美国中西部郊区的龙卷风到砸向意大利别墅屋顶的网球般大小的冰雹，温室气体排放造成的严峻天气正在打击世界上最重要的资产种类。'
)

# Phrase and Example
document.add_paragraph('Phrase: foot the bill')
document.add_paragraph('Explanation: 负担费用')
document.add_paragraph('Example: Once again it will be the taxpayer who has to foot the bill.')
document.add_paragraph('示例: 这一次掏腰包的又得是纳税人。')

# Insert an image
document.add_picture('F:/vscode_files/web技术/githubwork/web_program/TEXTREADING/TextImage/baciWelcome.jpg', width=Inches(4.0))

document.save('F:/vscode_files/web技术/githubwork/web_program/TEXTREADING/readingText/article.docx')
